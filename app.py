import os
import logging
from dotenv import load_dotenv
import anthropic
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def parse_ai_response(ai_response):
    # Define markers based on how sections are labeled in the AI response
    markers = ["===SUMMARY===", "===EXPERIENCE===", "===TECHNICAL SKILLS==="]
    sections = {"SUMMARY": "", "EXPERIENCE": "", "TECHNICAL SKILLS": ""}
    current_section = None
    for line in ai_response.split('\n'):
        if line in markers:
            current_section = line.replace('===', '').strip() # Extract section name
        elif current_section:
            sections[current_section] += line + '\n'
    return sections


def inject_custom_content(template_docx, ai_response):
    # Parse AI response into structured sections
    sections = parse_ai_response(ai_response)
    
    # Load the template document
    doc = Document(template_docx)
    
    # Define styles for section headers and content
    styles = doc.styles
    content_style = styles.add_style('Content', WD_STYLE_TYPE.PARAGRAPH)
    content_style.font.size = Pt(11)
    content_style.font.name = 'Arial'
    
    for paragraph in doc.paragraphs:
        if "<SUMMARY_PLACEHOLDER>" in paragraph.text:
            paragraph.text = paragraph.text.replace("<SUMMARY_PLACEHOLDER>", sections["SUMMARY"].strip())
            paragraph.style = content_style
        elif "<EXPERIENCE_PLACEHOLDER>" in paragraph.text:
            paragraph.text = paragraph.text.replace("<EXPERIENCE_PLACEHOLDER>", sections["EXPERIENCE"].strip())
            paragraph.style = content_style
        elif "<TECHNICAL_SKILLS_PLACEHOLDER>" in paragraph.text:
            paragraph.text = paragraph.text.replace("<TECHNICAL_SKILLS_PLACEHOLDER>", sections["TECHNICAL SKILLS"].strip())
            paragraph.style = content_style
    
    # Save the updated document as a Word file
    updated_docx_path = "Alessandro_Amenta_Resume.docx"
    doc.save(updated_docx_path)
    logging.info("Successfully updated the document with AI-generated content.")


def read_docx_file(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        content = '\n'.join(full_text)
        logging.info(f"Successfully read content from {file_path}")
        return content
    except Exception as e:
        logging.error(f"Error reading DOCX file content: {e}")
        return None

def read_text_file(file_path):
    try:
        with open(file_path, 'r') as file:
            content = file.read()
        logging.info(f"Successfully read content from {file_path}")
        return content
    except Exception as e:
        logging.error(f"Error reading text file content: {e}")
        return None

def write_docx_file(content, output_path="Alessandro_Amenta_Resume.docx"):
    try:
        doc = Document()
        for line in content.split('\n'):
            doc.add_paragraph(line)
        doc.save(output_path)
        logging.info(f"Successfully wrote custom resume to {output_path}")
    except Exception as e:
        logging.error(f"Error writing DOCX file: {e}")

def generate_custom_content(resume_content, job_description):
    try:
        load_dotenv()  # Load environment variables from .env file
        api_key = os.getenv('ANTHROPIC_API_KEY')
        if not api_key:
            logging.error("ANTHROPIC_API_KEY not found in .env file")
            return None

        client = anthropic.Anthropic(api_key=api_key)
        logging.info("Successfully created Anthropic client")

        system_prompt = ("You are an expert in early career professional resume crafting, ATS optimization, and recruitment. "
                         "Adjust the following resume so it best fits the job description provided.")

        prompt = (
        "Please customize and optimize the following resume to better align with the provided job description. "
        "Focus on enhancing the SUMMARY, EXPERIENCE, and TECHNICAL SKILLS sections to make the "
        "candidate more competitive for the role.\n\n"
        "When updating each section, keep the following guidelines in mind:\n\n"
        "1. Make the summary more concise, compelling and targeted towards the specific role. Highlight the "
        "candidate's most relevant skills and experiences. DO NOT directly mention the target company in the summary.\n"
        "2. For the EXPERIENCE section, keep all the original experiences and job titles intact. The number of bullet points for each experience should remain the same, and the overall length should be almost exactly the same as the original. "
        "However, you may slightly tweak the content of each bullet point to emphasize achievements and responsibilities that are most relevant to the target role. If needed, you can tweak the bulletpoint to showcaser relevant skills, but alwaysensure that the core experiences remain unchanged as well as the number and length of the bulletpoints, that should also remain the same overall.\n"
        "3. For the TECHNICAL SKILLS section, you may replace or update some of the skills to match the job requirements, "
        "but ensure the section length remains the same as the original.\n"
        "4. Do not make any changes to the other sections of the resume.\n\n"
        "Please output only the updated resume content (no other commentary) with each modified section clearly labeled as follows:\n\n"
        "===SUMMARY===\n"
        "[Insert updated summary here]\n\n"
        "===EXPERIENCE===\n"
        "[Insert updated experience section here]\n\n"
        "===TECHNICAL SKILLS===\n"
        "[Insert updated technical skills section here]\n\n"
        "Leave all other resume sections unchanged.\n\n"
        "Remember, the goal is to tailor the resume to the job description while maintaining the candidate's core experiences and overall identity. All original job titles and companies should remain the same. Ensure strict adherence of the above guidelines\n\n"
        "Here is the job description for reference:\n\n"
        f"{job_description}\n\n"
        "And here is the original resume content to be updated:\n\n"
        f"{resume_content}"
        )


        logging.info(f"Sending the following prompt to the AI: {prompt}")

        response = client.messages.create(
            model="claude-3-opus-20240229",
            max_tokens=4000,
            system=system_prompt,
            temperature=0.5,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": prompt
                        }
                    ]
                }
            ]
        )

        logging.info("Successfully generated custom resume content")
        logging.info(f"AI Response: {response.content[0].text.strip()}")
        return response.content[0].text.strip()
    except Exception as e:
        logging.error(f"Error generating custom resume: {e}")
        return None


if __name__ == "__main__":
    input_resume_path = "./context/resume.docx"
    job_description_path = "./context/jd.txt"
    output_resume_path = "./context/template.docx"
    
    job_description_content = read_text_file(job_description_path)
    resume_content = read_docx_file(input_resume_path)
    customized_content = generate_custom_content(resume_content, job_description_content)
    inject_custom_content(output_resume_path, customized_content)
